"""
Advanced document processing and ingestion system
"""
import os
import io
import re
import hashlib
import mimetypes
from typing import List, Dict, Any, Optional, Union, Generator, Tuple
from dataclasses import dataclass
from pathlib import Path
import logging
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
import asyncio
import aiofiles
from datetime import datetime

# Document processing imports
import PyPDF2
import docx
import pandas as pd
from bs4 import BeautifulSoup
import markdown
import json
import xml.etree.ElementTree as ET

# NLP imports
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
import spacy
from textstat import flesch_reading_ease, flesch_kincaid_grade

# ML imports
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
import numpy as np

from config import config, SUPPORTED_FORMATS
from vector_store import Document
from ml_utils import QueryAnalyzer, DocumentQuality

logger = logging.getLogger(__name__)

@dataclass
class ProcessingMetrics:
    """Document processing metrics"""
    total_documents: int = 0
    processed_documents: int = 0
    failed_documents: int = 0
    total_chunks: int = 0
    processing_time: float = 0.0
    average_chunk_size: float = 0.0

@dataclass
class DocumentMetadata:
    """Enhanced document metadata"""
    filename: str
    file_type: str
    file_size: int
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    author: Optional[str] = None
    title: Optional[str] = None
    language: str = "en"
    page_count: Optional[int] = None
    word_count: int = 0
    character_count: int = 0
    readability_score: Optional[float] = None
    topics: List[str] = None
    entities: List[str] = None
    quality_score: Optional[float] = None

    def __post_init__(self):
        if self.topics is None:
            self.topics = []
        if self.entities is None:
            self.entities = []

class DocumentProcessor:
    """Advanced document processing with ML-enhanced chunking and metadata extraction"""
    
    def __init__(self):
        self.chunk_size = config.vector_store.chunk_size
        self.chunk_overlap = config.vector_store.chunk_overlap
        self.supported_formats = SUPPORTED_FORMATS
        
        # Initialize NLP components
        self.query_analyzer = QueryAnalyzer()
        
        # Load spaCy model for advanced processing
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except:
            logger.warning("spaCy model not found. Some features will be disabled.")
            self.nlp = None
        
        # Initialize TF-IDF for topic extraction
        self.tfidf_vectorizer = TfidfVectorizer(
            max_features=100,
            stop_words='english',
            ngram_range=(1, 2)
        )
        
        # Document quality assessment
        self.quality_thresholds = {
            'min_length': 50,
            'min_readability': 10,
            'min_information_density': 0.3
        }
        
        # Processing statistics
        self.metrics = ProcessingMetrics()
    
    async def process_file(self, file_path: str, **kwargs) -> List[Document]:
        """Process a single file asynchronously"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Determine file type
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract content based on file type
            content = await self._extract_content(file_path)
            
            if not content or len(content.strip()) < self.quality_thresholds['min_length']:
                logger.warning(f"File content too short or empty: {file_path}")
                return []
            
            # Extract metadata
            metadata = await self._extract_metadata(file_path, content)
            
            # Intelligent chunking
            chunks = await self._intelligent_chunking(content, metadata)
            
            # Create documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc_id = f"{file_path.stem}_{i}"
                chunk_metadata = metadata.__dict__.copy()
                chunk_metadata.update({
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'source_file': str(file_path),
                    'chunk_method': 'intelligent'
                })
                
                document = Document(
                    id=doc_id,
                    content=chunk,
                    metadata=chunk_metadata
                )
                documents.append(document)
            
            self.metrics.processed_documents += 1
            self.metrics.total_chunks += len(chunks)
            
            logger.info(f"Processed {file_path.name}: {len(chunks)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            self.metrics.failed_documents += 1
            return []
    
    async def process_multiple_files(self, file_paths: List[str], 
                                   max_workers: int = 4) -> List[Document]:
        """Process multiple files concurrently"""
        start_time = datetime.now()
        self.metrics.total_documents = len(file_paths)
        
        # Process files concurrently
        tasks = [self.process_file(file_path) for file_path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Flatten results
        all_documents = []
        for result in results:
            if isinstance(result, list):
                all_documents.extend(result)
            elif isinstance(result, Exception):
                logger.error(f"Processing error: {result}")
        
        # Update metrics
        self.metrics.processing_time = (datetime.now() - start_time).total_seconds()
        if self.metrics.total_chunks > 0:
            total_chars = sum(len(doc.content) for doc in all_documents)
            self.metrics.average_chunk_size = total_chars / self.metrics.total_chunks
        
        logger.info(f"Processed {len(file_paths)} files: {len(all_documents)} total documents")
        return all_documents
    
    async def _extract_content(self, file_path: Path) -> str:
        """Extract content from file based on format"""
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return await self._extract_pdf_content(file_path)
            elif file_extension == '.docx':
                return await self._extract_docx_content(file_path)
            elif file_extension == '.txt':
                return await self._extract_text_content(file_path)
            elif file_extension == '.md':
                return await self._extract_markdown_content(file_path)
            elif file_extension == '.html':
                return await self._extract_html_content(file_path)
            elif file_extension == '.csv':
                return await self._extract_csv_content(file_path)
            elif file_extension == '.json':
                return await self._extract_json_content(file_path)
            elif file_extension == '.xml':
                return await self._extract_xml_content(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return ""
    
    async def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract content from PDF files"""
        try:
            async with aiofiles.open(file_path, 'rb') as file:
                pdf_data = await file.read()
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))
            content = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        content.append(f"[Page {page_num + 1}]\n{text}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1} from {file_path}: {e}")
            
            return "\n\n".join(content)
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return ""
    
    async def _extract_docx_content(self, file_path: Path) -> str:
        """Extract content from DOCX files"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = [cell.text.strip() for cell in row.cells]
                    table_content.append(" | ".join(row_content))
                if table_content:
                    content.append("\n".join(table_content))
            
            return "\n\n".join(content)
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return ""
    
    async def _extract_text_content(self, file_path: Path) -> str:
        """Extract content from text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            return content
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    content = await file.read()
                return content
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                return ""
    
    async def _extract_markdown_content(self, file_path: Path) -> str:
        """Extract content from Markdown files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                md_content = await file.read()
            
            # Convert markdown to text (remove formatting)
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
            
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {e}")
            return ""
    
    async def _extract_html_content(self, file_path: Path) -> str:
        """Extract content from HTML files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                html_content = await file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {e}")
            return ""
    
    async def _extract_csv_content(self, file_path: Path) -> str:
        """Extract content from CSV files"""
        try:
            df = pd.read_csv(file_path)
            
            # Convert to readable format
            content = []
            content.append(f"CSV file with {len(df)} rows and {len(df.columns)} columns")
            content.append(f"Columns: {', '.join(df.columns.tolist())}")
            
            # Add sample data
            if len(df) > 0:
                content.append("\nSample data:")
                content.append(df.head(5).to_string(index=False))
            
            # Add summary statistics for numeric columns
            numeric_columns = df.select_dtypes(include=[np.number]).columns
            if len(numeric_columns) > 0:
                content.append("\nNumeric column statistics:")
                content.append(df[numeric_columns].describe().to_string())
            
            return "\n\n".join(content)
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            return ""
    
    async def _extract_json_content(self, file_path: Path) -> str:
        """Extract content from JSON files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                json_content = await file.read()
            
            data = json.loads(json_content)
            
            # Convert JSON to readable text
            def json_to_text(obj, level=0):
                result = []
                indent = "  " * level
                
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            result.append(f"{indent}{key}:")
                            result.append(json_to_text(value, level + 1))
                        else:
                            result.append(f"{indent}{key}: {value}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        result.append(f"{indent}[{i}]:")
                        result.append(json_to_text(item, level + 1))
                else:
                    result.append(f"{indent}{obj}")
                
                return "\n".join(result)
            
            return json_to_text(data)
            
        except Exception as e:
            logger.error(f"Error processing JSON {file_path}: {e}")
            return ""
    
    async def _extract_xml_content(self, file_path: Path) -> str:
        """Extract content from XML files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                xml_content = await file.read()
            
            root = ET.fromstring(xml_content)
            
            # Extract all text content
            def extract_text(element, level=0):
                result = []
                indent = "  " * level
                
                if element.text and element.text.strip():
                    result.append(f"{indent}{element.tag}: {element.text.strip()}")
                elif element.tag:
                    result.append(f"{indent}{element.tag}:")
                
                for child in element:
                    result.append(extract_text(child, level + 1))
                
                return "\n".join(filter(None, result))
            
            return extract_text(root)
            
        except Exception as e:
            logger.error(f"Error processing XML {file_path}: {e}")
            return ""
    
    async def _extract_metadata(self, file_path: Path, content: str) -> DocumentMetadata:
        """Extract comprehensive metadata from document"""
        try:
            # Basic file metadata
            stat = file_path.stat()
            
            metadata = DocumentMetadata(
                filename=file_path.name,
                file_type=file_path.suffix.lower(),
                file_size=stat.st_size,
                created_date=datetime.fromtimestamp(stat.st_ctime),
                modified_date=datetime.fromtimestamp(stat.st_mtime),
                word_count=len(word_tokenize(content)),
                character_count=len(content)
            )
            
            # Extract title from content
            metadata.title = self._extract_title(content)
            
            # Calculate readability
            if len(content) > 100:
                try:
                    metadata.readability_score = flesch_reading_ease(content)
                except:
                    metadata.readability_score = 50.0
            
            # Extract topics using TF-IDF
            if len(content) > 200:
                metadata.topics = await self._extract_topics(content)
            
            # Extract entities using spaCy
            if self.nlp and len(content) > 100:
                metadata.entities = await self._extract_entities(content)
            
            # Calculate quality score
            metadata.quality_score = self._calculate_quality_score(content, metadata)
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {e}")
            return DocumentMetadata(
                filename=file_path.name,
                file_type=file_path.suffix.lower(),
                file_size=0,
                word_count=len(word_tokenize(content)) if content else 0,
                character_count=len(content) if content else 0
            )
    
    def _extract_title(self, content: str) -> Optional[str]:
        """Extract title from content"""
        lines = content.strip().split('\n')
        
        # Look for title patterns
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if len(line) > 10 and len(line) < 200:
                # Check if it looks like a title
                if (line.isupper() or 
                    line.count('.') == 0 or 
                    line.startswith('#')):
                    return line.replace('#', '').strip()
        
        # Fallback to first non-empty line
        for line in lines[:3]:
            line = line.strip()
            if len(line) > 10:
                return line[:100] + "..." if len(line) > 100 else line
        
        return None
    
    async def _extract_topics(self, content: str, max_topics: int = 5) -> List[str]:
        """Extract topics using TF-IDF"""
        try:
            # Preprocess text
            sentences = sent_tokenize(content)
            if len(sentences) < 2:
                return []
            
            # Fit TF-IDF
            tfidf_matrix = self.tfidf_vectorizer.fit_transform(sentences)
            feature_names = self.tfidf_vectorizer.get_feature_names_out()
            
            # Get top terms
            scores = np.asarray(tfidf_matrix.sum(axis=0)).flatten()
            top_indices = scores.argsort()[-max_topics:][::-1]
            
            topics = [feature_names[i] for i in top_indices if scores[i] > 0.1]
            return topics
            
        except Exception as e:
            logger.error(f"Error extracting topics: {e}")
            return []
    
    async def _extract_entities(self, content: str, max_entities: int = 10) -> List[str]:
        """Extract named entities using spaCy"""
        try:
            # Process text with spaCy
            doc = self.nlp(content[:10000])  # Limit for performance
            
            # Extract entities
            entities = []
            for ent in doc.ents:
                if ent.label_ in ["PERSON", "ORG", "GPE", "PRODUCT", "EVENT", "WORK_OF_ART"]:
                    entity_text = ent.text.strip()
                    if len(entity_text) > 2 and entity_text not in entities:
                        entities.append(entity_text)
            
            return entities[:max_entities]
            
        except Exception as e:
            logger.error(f"Error extracting entities: {e}")
            return []
    
    def _calculate_quality_score(self, content: str, metadata: DocumentMetadata) -> float:
        """Calculate document quality score"""
        try:
            score = 0.0
            max_score = 1.0
            
            # Length score (0.2 weight)
            length_score = min(1.0, len(content) / 1000)
            score += 0.2 * length_score
            
            # Readability score (0.3 weight)
            if metadata.readability_score:
                readability_score = min(1.0, max(0.0, metadata.readability_score / 100))
                score += 0.3 * readability_score
            else:
                score += 0.15  # Default middle score
            
            # Information density (0.2 weight)
            words = word_tokenize(content.lower())
            if len(words) > 0:
                unique_words = len(set(words))
                density = unique_words / len(words)
                score += 0.2 * min(1.0, density * 2)
            
            # Structure score (0.2 weight)
            structure_score = 0.0
            if metadata.title:
                structure_score += 0.5
            if len(sent_tokenize(content)) > 5:
                structure_score += 0.3
            if any(char in content for char in ['.', '!', '?']):
                structure_score += 0.2
            score += 0.2 * structure_score
            
            # Topic richness (0.1 weight)
            topic_score = min(1.0, len(metadata.topics) / 5) if metadata.topics else 0.5
            score += 0.1 * topic_score
            
            return min(1.0, score)
            
        except Exception as e:
            logger.error(f"Error calculating quality score: {e}")
            return 0.5
    
    async def _intelligent_chunking(self, content: str, metadata: DocumentMetadata) -> List[str]:
        """Intelligent content chunking using multiple strategies"""
        try:
            # Choose chunking strategy based on content type and quality
            if metadata.file_type in ['.pdf', '.docx']:
                return await self._semantic_chunking(content)
            elif metadata.file_type in ['.csv', '.json', '.xml']:
                return await self._structural_chunking(content)
            else:
                return await self._hybrid_chunking(content)
                
        except Exception as e:
            logger.error(f"Error in intelligent chunking: {e}")
            return self._simple_chunking(content)
    
    async def _semantic_chunking(self, content: str) -> List[str]:
        """Semantic-aware chunking for documents"""
        try:
            sentences = sent_tokenize(content)
            chunks = []
            current_chunk = []
            current_length = 0
            
            for sentence in sentences:
                sentence_length = len(sentence)
                
                # If adding this sentence would exceed chunk size
                if current_length + sentence_length > self.chunk_size and current_chunk:
                    # Create chunk with overlap
                    chunk_text = ' '.join(current_chunk)
                    chunks.append(chunk_text)
                    
                    # Start new chunk with overlap
                    overlap_sentences = current_chunk[-2:] if len(current_chunk) > 2 else current_chunk
                    current_chunk = overlap_sentences + [sentence]
                    current_length = sum(len(s) for s in current_chunk)
                else:
                    current_chunk.append(sentence)
                    current_length += sentence_length
            
            # Add final chunk
            if current_chunk:
                chunks.append(' '.join(current_chunk))
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error in semantic chunking: {e}")
            return self._simple_chunking(content)
    
    async def _structural_chunking(self, content: str) -> List[str]:
        """Structure-aware chunking for structured documents"""
        try:
            # Split by natural structural boundaries
            sections = re.split(r'\n\s*\n|\n\s*[-=]{3,}\s*\n', content)
            
            chunks = []
            for section in sections:
                section = section.strip()
                if len(section) > self.chunk_size:
                    # Further split large sections
                    sub_chunks = self._simple_chunking(section)
                    chunks.extend(sub_chunks)
                elif len(section) > 50:  # Skip very small sections
                    chunks.append(section)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error in structural chunking: {e}")
            return self._simple_chunking(content)
    
    async def _hybrid_chunking(self, content: str) -> List[str]:
        """Hybrid chunking combining multiple strategies"""
        try:
            # First try semantic chunking
            semantic_chunks = await self._semantic_chunking(content)
            
            # If chunks are too uneven, fallback to simple chunking
            if semantic_chunks:
                chunk_sizes = [len(chunk) for chunk in semantic_chunks]
                size_variance = np.var(chunk_sizes) if len(chunk_sizes) > 1 else 0
                
                if size_variance < (self.chunk_size * 0.5) ** 2:
                    return semantic_chunks
            
            return self._simple_chunking(content)
            
        except Exception as e:
            logger.error(f"Error in hybrid chunking: {e}")
            return self._simple_chunking(content)
    
    def _simple_chunking(self, content: str) -> List[str]:
        """Simple character-based chunking with sentence boundary awareness"""
        if len(content) <= self.chunk_size:
            return [content]
        
        chunks = []
        start = 0
        
        while start < len(content):
            end = start + self.chunk_size
            
            if end >= len(content):
                chunks.append(content[start:])
                break
            
            # Try to find sentence boundary near the end
            chunk_end = end
            for i in range(end - 50, end + 50):
                if i < len(content) and content[i] in '.!?':
                    chunk_end = i + 1
                    break
            
            chunks.append(content[start:chunk_end])
            start = chunk_end - self.chunk_overlap
        
        return [chunk for chunk in chunks if len(chunk.strip()) > 20]
    
    def get_processing_metrics(self) -> ProcessingMetrics:
        """Get current processing metrics"""
        return self.metrics
    
    def reset_metrics(self) -> None:
        """Reset processing metrics"""
        self.metrics = ProcessingMetrics()

# Factory function
def create_document_processor(**kwargs) -> DocumentProcessor:
    """Create document processor instance"""
    return DocumentProcessor(**kwargs)

# Export classes
__all__ = [
    'DocumentProcessor',
    'DocumentMetadata', 
    'ProcessingMetrics',
    'create_document_processor'
]   